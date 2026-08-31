import logging
from io import BytesIO
from urllib.parse import quote

import openpyxl
import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from exports.docx_exporter import (
    add_sources_to_docx,
)
from exports.excel_exporter import (
    add_sources_to_excel,
)
from exports.html_parser import (
    parse_html_table,
)
from schemas.export import ExportRequest


logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/export",
    tags=["export"],
)

@router.post("/docx")
def export_docx(
    request: ExportRequest,
) -> StreamingResponse:
    try:
        tables, sources_text = parse_html_table(request.content)

        if not tables:
            raise HTTPException(
                status_code=400,
                detail="В содержимом не найдены таблицы",
            )

        doc = DocxDocument()
        doc.add_heading(
            "Экспорт адаптационных мероприятий",
            level=1,
        )

        table_added = False

        for idx, table_data in enumerate(tables, start=1):
            if len(table_data) < 2:
                continue

            if table_added:
                doc.add_page_break()

            rows = len(table_data)
            cols = len(table_data[0])

            if cols == 0:
                continue

            word_table = doc.add_table(
                rows=rows,
                cols=cols,
            )
            word_table.style = "Table Grid"

            for row_index, row_data in enumerate(table_data):
                normalized_row = list(row_data[:cols])

                if len(normalized_row) < cols:
                    normalized_row.extend(
                        [""] * (cols - len(normalized_row))
                    )

                for column_index, cell_text in enumerate(
                    normalized_row
                ):
                    cell = word_table.cell(
                        row_index,
                        column_index,
                    )
                    cell.text = str(cell_text or "")

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(
                                10 if row_index == 0 else 9
                            )

                            if row_index == 0:
                                run.font.bold = True

            table_added = True

        if not table_added:
            raise HTTPException(
                status_code=400,
                detail="В содержимом нет таблиц с данными",
            )

        # Источники добавляются один раз после всех таблиц.
        if sources_text:
            add_sources_to_docx(
                doc,
                sources_text,
            )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = _ensure_extension(
            request.filename,
            ".docx",
        )

        encoded_filename = quote(filename)

        return StreamingResponse(
            buffer,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": (
                    "attachment; "
                    f"filename*=UTF-8''{encoded_filename}"
                )
            },
        )

    except HTTPException:
        raise

    except Exception as exc:
        logger.exception(
            "Ошибка экспорта документа DOCX"
        )
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка экспорта DOCX: {exc}",
        ) from exc


@router.post("/excel")
def export_excel(
    request: ExportRequest,
) -> StreamingResponse:
    try:
        tables, sources_text = parse_html_table(request.content)

        if not tables:
            raise HTTPException(
                status_code=400,
                detail="В содержимом не найдены таблицы",
            )

        buffer = BytesIO()
        sheet_created = False

        with pd.ExcelWriter(
            buffer,
            engine="openpyxl",
        ) as writer:
            for idx, table_data in enumerate(
                tables,
                start=1,
            ):
                if len(table_data) < 2:
                    continue

                headers = [
                    str(value or "")
                    for value in table_data[0]
                ]

                num_cols = len(headers)

                if num_cols == 0:
                    continue

                cleaned_rows: list[list[str]] = []

                for row in table_data[1:]:
                    normalized_row = [
                        str(value or "")
                        for value in row[:num_cols]
                    ]

                    if len(normalized_row) < num_cols:
                        normalized_row.extend(
                            [""] * (
                                num_cols
                                - len(normalized_row)
                            )
                        )

                    cleaned_rows.append(normalized_row)

                if not cleaned_rows:
                    continue

                dataframe = pd.DataFrame(
                    cleaned_rows,
                    columns=headers,
                )

                sheet_name = f"Table_{idx}"[:31]

                dataframe.to_excel(
                    writer,
                    sheet_name=sheet_name,
                    index=False,
                )

                sheet_created = True
                worksheet = writer.sheets[sheet_name]

                for cell in worksheet[1]:
                    if cell.value is not None:
                        cell.font = openpyxl.styles.Font(
                            bold=True
                        )

                if sources_text:
                    start_row = len(dataframe) + 3

                    add_sources_to_excel(
                        worksheet,
                        sources_text,
                        start_row,
                    )

            if not sheet_created:
                worksheet = writer.book.create_sheet(
                    title="No_Data"
                )
                worksheet["A1"] = (
                    "Нет данных для экспорта"
                )
                worksheet.sheet_state = "visible"

            if writer.book.worksheets:
                writer.book.active = 0
                writer.book.worksheets[
                    0
                ].sheet_state = "visible"

        buffer.seek(0)

        filename = _ensure_extension(
            request.filename,
            ".xlsx",
        )

        encoded_filename = quote(filename)

        return StreamingResponse(
            buffer,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            headers={
                "Content-Disposition": (
                    "attachment; "
                    f"filename*=UTF-8''{encoded_filename}"
                )
            },
        )

    except HTTPException:
        raise

    except Exception as exc:
        logger.exception(
            "Ошибка экспорта файла Excel"
        )
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка экспорта Excel: {exc}",
        ) from exc


def _ensure_extension(
    filename: str,
    extension: str,
) -> str:
    normalized_filename = filename.strip()

    if not normalized_filename:
        normalized_filename = "export"

    if not normalized_filename.lower().endswith(extension):
        normalized_filename = (
            f"{normalized_filename}{extension}"
        )

    return normalized_filename
