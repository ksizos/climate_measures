import re

from docx.document import Document


def add_sources_to_docx(
    doc: Document,
    sources_text: str | None,
) -> None:
    """
    Добавляет раздел с опорными источниками в DOCX-документ.
    """
    if not sources_text or not sources_text.strip():
        return

    doc.add_paragraph()

    lines = sources_text.strip().splitlines()
    header_added = False

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            continue

        if "опорные источники" in line.lower():
            if not header_added:
                doc.add_heading(
                    "Опорные источники:",
                    level=3,
                )
                header_added = True

            line = re.sub(
                r"опорные источники[:\s]*",
                "",
                line,
                flags=re.IGNORECASE,
            ).strip()

            if not line:
                continue

        paragraph = doc.add_paragraph()

        url_match = re.search(
            r"https?://[^\s\]\[•\n]+",
            line,
        )

        if not url_match:
            paragraph.add_run(line)
            continue

        url = url_match.group(0)

        text_before, _, text_after = line.partition(url)

        if text_before.strip():
            paragraph.add_run(
                text_before.strip() + " "
            )

        url_run = paragraph.add_run(url)
        url_run.font.underline = True

        if text_after.strip():
            paragraph.add_run(
                " " + text_after.strip()
            )
